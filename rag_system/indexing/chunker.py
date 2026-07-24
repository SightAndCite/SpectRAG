from __future__ import annotations
import hashlib
import logging
import re
from pathlib import Path
from typing import Callable

import numpy as np
import tiktoken
# NOTE: faiss must be initialized before langchain_text_splitters. langchain-core's
# native deps (pydantic-core/orjson) load a runtime that segfaults faiss's
# BLAS/OpenMP if faiss initializes afterwards (macOS x86). Importing faiss here
# first — this module is the only importer of langchain_text_splitters — pins the
# safe order process-wide. Do not reorder these two imports.
import faiss  # noqa: F401  (import-order guard, see above)
from langchain_text_splitters import RecursiveCharacterTextSplitter

from rag_system.models import Chunk
from rag_system.language.detector import detect_language
from config import IndexingConfig, LanguageConfig

_SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?。！？])\s+")

logger = logging.getLogger(__name__)

# Regex patterns (module-level for single compilation)
_HEADING_RE          = re.compile(r"^(#{1,6})\s+(.+?)\s*#*$")           # markdown ATX heading
_NUMBERED_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)\s+(\S[^\n]{3,})$")  # "1.2 Background"
_CITE_BRACKET_RE     = re.compile(r"\[[\d,\s–\-]+\]")
_CITE_AUTHOR_RE      = re.compile(r"\(\w[^)]{2,40}\d{4}\w?\)")
_DOI_RE              = re.compile(r"doi:\s*10\.\d{4,}/\S+", re.IGNORECASE)
_URL_RE              = re.compile(r"https?://\S+")
# Matches 3+ single characters each separated by exactly one space:  "H E L L O"
# Boundary guards ensure we don't match chars that are mid-word (e.g. "am a s" in "I am a student")
_CHAR_SPACED_RE      = re.compile(r"(?<!\S)[\w.]( [\w.]){2,}(?!\S)", re.UNICODE)

# Sentence/structure separators for the recursive splitter, coarse → fine. This
# keeps splits on paragraph, then line, then sentence, then clause, then word
# boundaries (incl. CJK sentence marks) so chunks stay semantically coherent.
_SEPARATORS = ["\n\n", "\n", "。", "！", "？", ". ", "! ", "? ", "; ", ", ", " ", ""]

# tiktoken encoders are cached per encoding name (thread-safe, cheap to reuse).
_ENCODERS: dict[str, "tiktoken.Encoding"] = {}


def _get_encoder(name: str) -> "tiktoken.Encoding":
    enc = _ENCODERS.get(name)
    if enc is None:
        enc = tiktoken.get_encoding(name)
        _ENCODERS[name] = enc
    return enc


# Per-format page readers
# Each reader returns a list of (page_text, extra_meta_dict) tuples.
ReaderFn = Callable[[Path], list[tuple[str, dict]]]
_READERS: dict[str, ReaderFn] = {}


def _register(*exts: str) -> Callable[[ReaderFn], ReaderFn]:
    def decorator(fn: ReaderFn) -> ReaderFn:
        for ext in exts:
            _READERS[ext] = fn
        return fn
    return decorator


def _fix_char_spacing(text: str) -> str:
    """Collapse PDFs exported with per-character spacing: 'H E L L O' → 'HELLO'."""
    return _CHAR_SPACED_RE.sub(lambda m: m.group().replace(" ", ""), text)


@_register(".pdf")
def _read_pdf(path: Path) -> list[tuple[str, dict]]:
    # pdfplumber preserves reading order / layout far better than pypdf and can
    # recover tables; we keep per-page granularity so page numbers stay accurate.
    import pdfplumber
    pages: list[tuple[str, dict]] = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            pages.append((_fix_char_spacing(text), {"page": i + 1}))
    return pages


@_register(".txt")
def _read_txt(path: Path) -> list[tuple[str, dict]]:
    return [(path.read_text(encoding="utf-8", errors="replace"), {})]


@_register(".md")
def _read_md(path: Path) -> list[tuple[str, dict]]:
    return [(path.read_text(encoding="utf-8", errors="replace"), {})]


@_register(".html", ".htm")
def _read_html(path: Path) -> list[tuple[str, dict]]:
    from bs4 import BeautifulSoup
    html = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return [(soup.get_text(separator="\n"), {})]


def _docx_prose(doc) -> str:
    # Use Word's real heading STYLES (not text/regex guessing): a paragraph styled
    # "Heading N"/"Title" is re-emitted as a markdown heading so the downstream
    # heading-stack segmentation gives it a proper section path. (doc.paragraphs
    # excludes table cell text, so tables are handled separately without dupes.)
    lines: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        s = ((p.style.name if p.style else "") or "").lower()
        if s.startswith("heading"):
            try:
                level = int(s.split()[-1])
            except ValueError:
                level = 1
            lines.append("#" * min(max(level, 1), 6) + " " + t)
        elif s == "title":
            lines.append("# " + t)
        else:
            lines.append(t)
    return "\n".join(lines)


@_register(".docx")
def _read_docx(path: Path) -> list[tuple[str, dict]]:
    import docx
    return [(_docx_prose(docx.Document(str(path))), {})]


def _table_to_md(rows: list) -> str:
    """Render extracted table rows as pipe-separated markdown lines."""
    clean = [
        [("" if c is None else str(c)).replace("\n", " ").strip() for c in (r or [])]
        for r in (rows or [])
    ]
    clean = [r for r in clean if any(cell for cell in r)]
    return "\n".join(" | ".join(r) for r in clean)


# Helpers

def _extract_references(text: str) -> list[str]:
    refs: list[str] = []
    refs.extend(_CITE_BRACKET_RE.findall(text))
    refs.extend(_CITE_AUTHOR_RE.findall(text))
    refs.extend(_DOI_RE.findall(text))
    refs.extend(m.rstrip(".,;)") for m in _URL_RE.findall(text))
    return list(dict.fromkeys(refs))


def _heading_of(line: str) -> tuple[int, str] | None:
    """Return (level, title) if the line is a heading, else None."""
    line = line.strip()
    if not line:
        return None
    m = _HEADING_RE.match(line)
    if m:
        return len(m.group(1)), m.group(2).strip()
    m = _NUMBERED_HEADING_RE.match(line)
    if m:
        # depth = number of dotted components: "1" → 1, "1.2" → 2, "1.2.3" → 3
        return m.group(1).count(".") + 1, m.group(0).strip()
    return None


def _segment_by_heading(
    text: str, stack: list[tuple[int, str]]
) -> tuple[list[tuple[list[str], str]], list[tuple[int, str]]]:
    """Split text into (section_path, block_text) segments using a heading stack.

    The `stack` (level, title) is passed in and returned so it PERSISTS across
    pages of the same document: a heading opened on page 1 still scopes the body
    text that continues onto page 2, instead of the section context resetting at
    every page boundary. Each segment's `section_path` is the active heading trail
    (deepest 3 levels); block boundaries fall on headings so a chunk never
    straddles two sections.
    """
    body: list[str] = []
    segments: list[tuple[list[str], str]] = []

    def path() -> list[str]:
        return [t for _lvl, t in stack][-3:]

    def flush() -> None:
        block = "\n".join(body).strip()
        if block:
            segments.append((path(), block))
        body.clear()

    for line in text.splitlines():
        h = _heading_of(line)
        if h is None:
            body.append(line)
            continue
        # New heading: close the previous section, then update the stack.
        flush()
        level, title = h
        while stack and stack[-1][0] >= level:
            stack.pop()
        stack.append((level, title))
        body.append(line)  # keep the heading text inside its own section's chunk

    flush()
    return segments, stack


# Main class

class DocumentChunker:
    """
    Splits documents into structure- and token-aware overlapping chunks.

    Per document a strategy is chosen (config `chunking_strategy`, default "auto"):
      • structural — heading-aware segmentation (heading stack persists across
        pages) + recursive paragraph/sentence/word splitting, real tiktoken sizing
      • semantic  — for heading-less prose, split where adjacent sentences are
        semantically distant (needs an embedder), then enforce the token budget
    Common to both: undersized tail chunks merge into a neighbour, and a
    sentence-bounded overlap tail bridges section/page boundaries. Tables in
    PDF/DOCX are extracted as dedicated row-boundary-aware chunks (never chopped),
    and excluded from the prose stream to avoid duplication.
    Supports PDF, TXT, MD, HTML, and DOCX via the _READERS registry.
    """

    def __init__(
        self,
        cfg: IndexingConfig,
        lang_cfg: LanguageConfig | None = None,
        embedder=None,
    ) -> None:
        self._chunk_size    = cfg.chunk_size
        self._chunk_overlap = cfg.chunk_overlap
        self._min_tokens    = cfg.min_chunk_tokens
        self._hash_prefix   = cfg.chunk_id_hash_prefix
        self._strategy      = getattr(cfg, "chunking_strategy", "auto")
        self._tables        = getattr(cfg, "table_extraction", True)
        self._sem_pct       = getattr(cfg, "semantic_breakpoint_percentile", 90.0)
        self._lang_cfg      = lang_cfg
        self._embedder      = embedder   # enables the semantic strategy; optional
        self._encoder       = _get_encoder(cfg.tiktoken_encoding)
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=cfg.chunk_size,
            chunk_overlap=cfg.chunk_overlap,
            length_function=self._token_len,
            separators=_SEPARATORS,
            keep_separator=True,
        )

    def _token_len(self, text: str) -> int:
        return len(self._encoder.encode(text, disallowed_special=()))

    def _sentence_tail(self, text: str, max_tokens: int) -> str:
        """Last whole sentence(s) of `text`, up to `max_tokens` — used as the
        overlap carried into the next chunk across a section/page boundary."""
        if max_tokens <= 0 or not text.strip():
            return ""
        parts = re.split(r"(?<=[.!?。！？])\s+", text.strip())
        tail: list[str] = []
        total = 0
        for s in reversed(parts):
            t = self._token_len(s)
            if tail and total + t > max_tokens:
                break
            tail.insert(0, s)
            total += t
        out = " ".join(tail).strip()
        # A single sentence longer than the budget: hard-trim to the last tokens.
        if self._token_len(out) > max_tokens * 2:
            ids = self._encoder.encode(out, disallowed_special=())[-max_tokens:]
            out = self._encoder.decode(ids)
        return out

    def _merge_small(self, pieces: list[str]) -> list[str]:
        """Fold chunks below the minimum token size into an adjacent chunk.

        Undersized pieces merge backward into the previous chunk; a small leading
        piece (e.g. a lone heading) has no previous, so it merges forward instead.
        """
        merged: list[str] = []
        for p in pieces:
            p = p.strip()
            if not p:
                continue
            if merged and self._token_len(p) < self._min_tokens:
                merged[-1] = f"{merged[-1]} {p}".strip()
            else:
                merged.append(p)
        if len(merged) > 1 and self._token_len(merged[0]) < self._min_tokens:
            merged[1] = f"{merged[0]} {merged[1]}".strip()
            merged.pop(0)
        return merged

    def _detect_lang(self, text: str) -> str:
        lc = self._lang_cfg
        fb = lc.fallback_language if lc else "en"
        return detect_language(text, fb) if (lc and lc.auto_detect) else fb

    def _make_chunk(self, text: str, doc_id: str, position: int, meta: dict,
                    section_path: list[str], lang: str) -> Chunk:
        # Include the source path so two files with the same stem (doc_id defaults
        # to path.stem) can't collide into one chunk_id and silently drop a chunk.
        chunk_id = hashlib.sha256(
            f"{meta.get('source', '')}:{doc_id}:{position}:{text[:self._hash_prefix]}".encode()
        ).hexdigest()[:16]
        return Chunk(
            chunk_id=chunk_id, doc_id=doc_id, text=text, position=position,
            section_path=section_path, references=_extract_references(text),
            language=lang, metadata={**meta, "language": lang},
        )

    def _select_strategy(self, text: str) -> str:
        """auto → structural when the doc has headings, else semantic (if an
        embedder is available) for unstructured prose."""
        if self._strategy in ("structural", "semantic"):
            return self._strategy
        headings = sum(1 for ln in text.splitlines() if _heading_of(ln))
        if headings >= 2:
            return "structural"
        return "semantic" if self._embedder is not None else "structural"

    def _read_with_tables(self, path: Path, suffix: str) -> list[tuple[str, dict, str]]:
        """Read a PDF/DOCX as ordered (text, meta, kind) items where kind is
        'prose' or 'table'. Table regions are excluded from the prose so a table
        is never both flowing-text and a table chunk (no duplication)."""
        items: list[tuple[str, dict, str]] = []
        if suffix == ".pdf":
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages):
                    meta = {"page": i + 1}
                    tables = page.find_tables()
                    if tables:
                        boxes = [t.bbox for t in tables]

                        def _outside(obj, boxes=boxes) -> bool:
                            cx = (obj["x0"] + obj["x1"]) / 2
                            cy = (obj["top"] + obj["bottom"]) / 2
                            return not any(
                                b[0] <= cx <= b[2] and b[1] <= cy <= b[3] for b in boxes
                            )

                        prose = page.filter(_outside).extract_text() or ""
                        for t in tables:
                            md = _table_to_md(t.extract())
                            if md:
                                items.append((md, meta, "table"))
                    else:
                        prose = page.extract_text() or ""
                    items.append((_fix_char_spacing(prose), meta, "prose"))
        else:  # .docx
            import docx
            d = docx.Document(str(path))
            items.append((_docx_prose(d), {}, "prose"))
            for t in d.tables:
                md = _table_to_md([[c.text for c in r.cells] for r in t.rows])
                if md:
                    items.append((md, {}, "table"))
        return items

    def _split_table(self, md: str) -> list[str]:
        """Split an oversized table on ROW boundaries (repeating the header row),
        so a table is never chopped mid-row."""
        if self._token_len(md) <= self._chunk_size:
            return [md]
        rows = md.split("\n")
        header = rows[0] if rows else ""
        out: list[str] = []
        cur: list[str] = []
        cur_tok = 0
        for r in rows:
            rt = self._token_len(r)
            if cur and cur_tok + rt > self._chunk_size:
                out.append("\n".join(cur))
                cur = [header] if header else []
                cur_tok = self._token_len(header) if cur else 0
            cur.append(r)
            cur_tok += rt
        if cur:
            out.append("\n".join(cur))
        return out

    def _chunk_semantic(self, text: str) -> list[str]:
        """Embedding-based splitting for heading-less prose: break where adjacent
        sentences are semantically distant (distance ≥ the configured percentile),
        then enforce the token budget with the recursive splitter."""
        sents = [s for s in _SENTENCE_SPLIT_RE.split(text.strip()) if s.strip()]
        if self._embedder is None or len(sents) < 4:
            return self._merge_small(self._splitter.split_text(text))
        embs = np.asarray(self._embedder.embed(sents, kind="document"), dtype=np.float32)
        dists = [1.0 - float(np.dot(embs[i], embs[i + 1])) for i in range(len(embs) - 1)]
        thr = float(np.percentile(dists, self._sem_pct)) if dists else 1.0
        groups: list[str] = []
        cur = [sents[0]]
        for i, d in enumerate(dists):
            if d >= thr and self._token_len(" ".join(cur)) >= self._min_tokens:
                groups.append(" ".join(cur))
                cur = [sents[i + 1]]
            else:
                cur.append(sents[i + 1])
        groups.append(" ".join(cur))
        out: list[str] = []
        for g in groups:
            if self._token_len(g) > self._chunk_size:
                out.extend(self._splitter.split_text(g))
            else:
                out.append(g)
        return self._merge_small(out)

    def chunk_file(self, path: Path | str, doc_id: str | None = None) -> list[Chunk]:
        path = Path(path)
        suffix = path.suffix.lower()
        reader = _READERS.get(suffix)
        if reader is None:
            logger.warning("No reader for '%s', falling back to plain text", suffix)
            reader = _READERS[".txt"]
            suffix = ".txt"

        doc_id = doc_id or path.stem
        source = str(path)

        # Read as ordered (text, meta, kind) items. Tables are pulled out as their
        # own items for PDF/DOCX so prose-splitting never chops them.
        if self._tables and suffix in (".pdf", ".docx"):
            items = self._read_with_tables(path, suffix)
        else:
            items = [(t, m, "prose") for t, m in reader(path)]

        prose_text = "\n".join(t for t, _m, kind in items if kind == "prose")
        strategy = self._select_strategy(prose_text)

        chunks: list[Chunk] = []
        stack: list[tuple[int, str]] = []   # heading stack — persists across pages
        prev_tail = ""                       # overlap carried across boundaries

        for text, meta, kind in items:
            meta = {**meta, "source": source}
            if kind == "table":
                prev_tail = ""  # a table breaks the prose flow — don't bridge overlap
                lang = self._detect_lang(text)
                for tt in self._split_table(text):
                    chunks.append(self._make_chunk(
                        tt, doc_id, len(chunks),
                        {**meta, "content_type": "table"}, ["(table)"], lang))
                continue
            if not text.strip():
                continue
            if strategy == "semantic":
                lang = self._detect_lang(text)
                for pi, piece in enumerate(self._chunk_semantic(text)):
                    if pi == 0 and prev_tail:
                        piece = f"{prev_tail}\n{piece}".strip()
                    chunks.append(self._make_chunk(
                        piece, doc_id, len(chunks), meta, [], lang))
                    prev_tail = self._sentence_tail(piece, self._chunk_overlap)
            else:  # structural (heading-aware)
                page_chunks, stack, prev_tail = self._split_text(
                    text, doc_id, len(chunks), meta, stack, prev_tail)
                chunks.extend(page_chunks)

        logger.debug("Chunked '%s' → %d chunks (%s)", path.name, len(chunks), strategy)
        return chunks

    def _split_text(
        self,
        text: str,
        doc_id: str,
        start_position: int,
        meta: dict,
        stack: list[tuple[int, str]],
        prev_tail: str,
    ) -> tuple[list[Chunk], list[tuple[int, str]], str]:
        """Structural strategy: heading-aware segmentation + recursive splitting,
        with the heading stack and overlap tail threaded across pages."""
        if not text.strip():
            return [], stack, prev_tail

        lang = self._detect_lang(text)
        chunks: list[Chunk] = []
        segments, stack = _segment_by_heading(text, stack)
        for section_path, block in segments:
            pieces = self._merge_small(self._splitter.split_text(block))
            for pi, piece in enumerate(pieces):
                # Bridge overlap across the section/page boundary: the first piece
                # of each block gets the tail of the previous chunk prepended. (In-
                # block pieces already overlap via the recursive splitter.)
                if pi == 0 and prev_tail:
                    piece = f"{prev_tail}\n{piece}".strip()
                chunks.append(self._make_chunk(
                    piece, doc_id, start_position + len(chunks),
                    meta, section_path, lang))
                prev_tail = self._sentence_tail(piece, self._chunk_overlap)
        return chunks, stack, prev_tail
